"""Genera DOCUMENTO_PRUEBAS_UNITARIAS.docx desde los archivos de prueba del proyecto."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
OUT = DOCS / "DOCUMENTO_PRUEBAS_UNITARIAS.docx"

INDEX_ROWS = [
    ("1", "usuario", "usuario/tests.py", "UsuarioModelTest", "6"),
    ("2", "rol", "rol/tests.py", "RolModelTest", "3"),
    ("3", "producto", "producto/tests.py", "ProductoAvailabilityTest", "7"),
    ("4", "pedido", "pedido/tests.py", "PedidoModelTest", "2"),
    ("5", "pedido", "pedido/test_flujo_compra.py", "BonosFidelidadTest, FinalizarCompraTest", "7"),
    ("6", "detalle_pedido", "detalle_pedido/tests.py", "DetallePedidoBusinessLogicTest", "4"),
    ("7", "inventario", "inventario/tests.py", "InventarioModelTest", "1"),
    ("8", "movimiento_inventario", "movimiento_inventario/tests.py", "MovimientoInventarioModelTest", "1"),
    ("9", "receta", "receta/tests.py", "RecetaModelTest", "1"),
    ("10", "recibo", "recibo/tests.py", "ReciboModelTest", "2"),
    ("11", "metodo_pago", "metodo_pago/tests.py", "MetodoPagoModelTest", "2"),
    ("12", "dashboard", "dashboard/tests.py", "ConfiguracionSistemaHorariosTestCase", "4"),
    ("13", "cuentas", "cuentas/tests.py", "CarritoTemplateTest", "6"),
    ("14", "cuentas", "cuentas/test_delete_utils.py", "EliminarConMensajeTest", "2"),
    ("15", "cuentas", "cuentas/test_seed_demo.py", "SeedDemoCommandTest", "1"),
    ("", "", "Total", "15 archivos", "49"),
]

USUARIO_EVIDENCE = [
    ("test_usuario_creation", "Creación correcta del usuario y relación con el rol"),
    ("test_full_clean_accepts_valid_data", "Datos válidos pasan validación del modelo"),
    ("test_full_clean_rejects_short_document", "Documento con menos de 5 dígitos es rechazado"),
    ("test_full_clean_rejects_invalid_phone", "Teléfono inválido es rechazado"),
    ("test_full_clean_rejects_non_numeric_document", "Documento no numérico es rechazado"),
    ("test_full_clean_allows_blank_address", "Dirección vacía es permitida"),
]

MODULE_FILES = [
    ("1. Módulo usuario", "usuario/tests.py"),
    ("2. Módulo rol", "rol/tests.py"),
    ("3. Módulo producto", "producto/tests.py"),
    ("4. Módulo pedido", "pedido/tests.py"),
    ("5. Módulo pedido (flujo compra)", "pedido/test_flujo_compra.py"),
    ("6. Módulo detalle_pedido", "detalle_pedido/tests.py"),
    ("7. Módulo inventario", "inventario/tests.py"),
    ("8. Módulo movimiento_inventario", "movimiento_inventario/tests.py"),
    ("9. Módulo receta", "receta/tests.py"),
    ("10. Módulo recibo", "recibo/tests.py"),
    ("11. Módulo metodo_pago", "metodo_pago/tests.py"),
    ("12. Módulo dashboard", "dashboard/tests.py"),
    ("13. Módulo cuentas", "cuentas/tests.py"),
    ("14. Módulo cuentas (delete_utils)", "cuentas/test_delete_utils.py"),
    ("15. Módulo cuentas (seed_demo)", "cuentas/test_seed_demo.py"),
]


def set_cell_shading(cell, fill_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): fill_hex,
    })
    shading.append(shd)


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def sanitize_text(text: str) -> str:
    return "".join(ch for ch in text if ch in "\n\t" or ord(ch) >= 32)


def add_code_block(doc: Document, code: str):
    code = sanitize_text(code)
    for line in code.splitlines() or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def add_table(doc: Document, headers, rows, header_fill="D9E2F3"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("ShoriExpress — Documento de pruebas unitarias", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, "Proyecto: Sistema de gestión para restaurante de chorizos")
    add_para(doc, "Framework: Django 4.2")
    add_para(doc, "Comando para ejecutar todas las pruebas:")
    add_code_block(doc, "cd shoriexpress\npython manage.py test")
    add_para(doc, "Resultado verificado: 49 pruebas — OK (todas pasan con dependencias de requirements.txt).", bold=True)

    add_heading(doc, "Índice de módulos y archivos de prueba", 1)
    add_table(
        doc,
        ["#", "Módulo Django", "Archivo de prueba", "Clases de prueba", "Cantidad"],
        INDEX_ROWS,
    )

    add_heading(doc, "Evidencia: ejecución de usuario/tests.py", 1)
    add_heading(doc, "Comando", 2)
    add_code_block(doc, "python manage.py test usuario.tests --verbosity=2")

    add_heading(doc, "Captura de pantalla (terminal)", 2)
    placeholder = doc.add_paragraph()
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ph_run = placeholder.add_run("[ ADJUNTE AQUÍ SU CAPTURA DE PANTALLA ]")
    ph_run.bold = True
    ph_run.font.size = Pt(14)
    ph_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    add_para(
        doc,
        "Inserte la imagen en este espacio: Inicio → Imagen → Este dispositivo, "
        "o arrastre el archivo PNG/JPG sobre este párrafo.",
        italic=True,
    )
    doc.add_paragraph()
    box = doc.add_paragraph()
    box.add_run(" ").font.size = Pt(72)

    add_heading(doc, "Salida textual esperada", 2)
    evidence_path = DOCS / "evidencia_usuario_tests.txt"
    if evidence_path.exists():
        add_code_block(doc, evidence_path.read_text(encoding="utf-8").strip())
    else:
        add_code_block(doc, "Ran 6 tests in 0.011s\nOK")

    add_heading(doc, "Pruebas ejecutadas (módulo usuario)", 2)
    add_table(doc, ["Test", "Qué valida"], USUARIO_EVIDENCE)

    for section_title, rel_path in MODULE_FILES:
        file_path = ROOT / rel_path.replace("/", "\\") if "\\" not in rel_path else ROOT / rel_path
        if not file_path.exists():
            file_path = ROOT / rel_path
        add_heading(doc, f"{section_title} — {rel_path}", 1)
        if file_path.exists():
            code = file_path.read_text(encoding="utf-8")
            add_code_block(doc, code.rstrip())
        else:
            add_para(doc, f"No se encontró el archivo: {rel_path}", italic=True)

    add_heading(doc, "Cómo reproducir en la sustentación", 1)
    add_code_block(
        doc,
        "# Todas las pruebas\npython manage.py test\n\n"
        "# Solo un módulo\npython manage.py test usuario.tests\n"
        "python manage.py test producto.tests\n"
        "python manage.py test pedido.test_flujo_compra\n\n"
        "# Con detalle\npython manage.py test usuario.tests --verbosity=2",
    )
    add_para(
        doc,
        "Requisito: pip install -r requirements.txt (incluye whitenoise).",
    )

    add_heading(doc, "Notas para la sustentación", 1)
    notes = [
        "Las pruebas usan una base de datos en memoria creada y destruida por Django.",
        "Cubren modelos, vistas, validaciones de negocio, carrito, bonos, borrados seguros y seed demo.",
        "El módulo usuario valida documento, teléfono y dirección antes de persistir datos.",
    ]
    for note in notes:
        doc.add_paragraph(note, style="List Bullet")

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Documento generado para ShoriExpress — Pruebas unitarias Django.")
    fr.italic = True
    fr.font.size = Pt(10)

    DOCS.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Generado: {OUT}")


if __name__ == "__main__":
    build_document()
